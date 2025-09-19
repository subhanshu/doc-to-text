"""
Enhanced FastAPI service with progress tracking for bulk operations
"""

import os
import shutil
import tempfile
import logging
import asyncio
import uuid
from concurrent.futures import ProcessPoolExecutor, TimeoutError as FuturesTimeout
from typing import List, Tuple, Dict, Any
from datetime import datetime

from fastapi import FastAPI, File, UploadFile, HTTPException, status, BackgroundTasks
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from pydantic_settings import BaseSettings

# Extraction libraries
import pdfplumber
import docx

# ---- Settings ----
class Settings(BaseSettings):
    MAX_FILE_SIZE_BYTES: int = 20 * 1024 * 1024  # 20 MB per file
    MAX_TOTAL_FILES: int = 8
    # Number of processes for ProcessPoolExecutor. Keep <= number of CPU cores.
    PROCESS_POOL_WORKERS: int = max(1, (os.cpu_count() or 1) - 1)
    # Limit concurrent requests to avoid OOM/CPU overcommit on small VM.
    MAX_CONCURRENT_REQUESTS: int = 3
    # Per-file extraction timeout (seconds)
    PER_FILE_TIMEOUT: int = 30
    # Overall request timeout
    REQUEST_TIMEOUT: int = 120
    # Temp dir base
    TEMP_DIR: str | None = None
    # Session retention period (seconds) - keep completed sessions for this long
    SESSION_RETENTION_SECONDS: int = 300  # 5 minutes
    # Maximum number of sessions to keep in memory
    MAX_SESSIONS: int = 100

    class Config:
        env_prefix = "EXTRACT_"

settings = Settings()

# Fix sensible defaults for a 4-core, 8GB VM if automatic detection returned large counts
if (os.cpu_count() or 1) >= 4:
    # leave one core for the event loop and webserver
    settings.PROCESS_POOL_WORKERS = min(settings.PROCESS_POOL_WORKERS, 3)
    settings.MAX_CONCURRENT_REQUESTS = min(settings.MAX_CONCURRENT_REQUESTS, 3)

# ---- Logging ----
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("extractor")

# ---- Pydantic Models for API Documentation ----

class FileResult(BaseModel):
    """Result of a successfully processed file"""
    filename: str
    text: str

class FileError(BaseModel):
    """Error information for a failed file processing"""
    filename: str
    error: str

class ProgressStatus(BaseModel):
    """Progress status for a processing session"""
    session_id: str
    status: str  # 'processing', 'completed', 'failed'
    total_files: int
    processed_files: int
    current_file: str | None = None
    results: List[FileResult] = []
    errors: List[FileError] = []
    start_time: datetime
    end_time: datetime | None = None

class ProgressResponse(BaseModel):
    """Response model for progress endpoint"""
    session_id: str
    status: str
    progress: float
    total_files: int
    processed_files: int
    current_file: str | None = None
    results: List[FileResult]
    errors: List[FileError]
    elapsed_time: float
    start_time: str
    end_time: str | None = None
    time_until_expiry: float | None = None

class SessionInfo(BaseModel):
    """Session information for listing endpoint"""
    session_id: str
    status: str
    progress: float
    total_files: int
    processed_files: int
    current_file: str | None = None
    elapsed_time: float
    start_time: str
    end_time: str | None = None
    time_until_expiry: float | None = None

class SessionsListResponse(BaseModel):
    """Response model for sessions list endpoint"""
    total_sessions: int
    sessions: List[SessionInfo]

class ExtractResponse(BaseModel):
    """Response model for extract endpoint"""
    results: List[FileResult]
    errors: List[FileError]

class ExtractProgressResponse(BaseModel):
    """Response model for extract-progress endpoint"""
    session_id: str
    status: str
    message: str

class APIInfoResponse(BaseModel):
    """Response model for root endpoint"""
    title: str
    version: str
    description: str
    endpoints: Dict[str, str]
    supported_formats: List[str]
    max_file_size: str
    max_files_per_request: int

class HealthResponse(BaseModel):
    """Response model for health endpoint"""
    status: str
    pool_workers: int

class MessageResponse(BaseModel):
    """Response model for simple message endpoints"""
    message: str

# In-memory store for progress (use Redis in production)
progress_store: Dict[str, ProgressStatus] = {}

# ---- FastAPI app ----
app = FastAPI(
    title="Document Text Extractor API",
    description="""
    ## Document Text Extractor API

    A production-ready FastAPI service for extracting text from PDF and Word (.docx) documents with real-time progress tracking.

    ### Features

    * **Multi-format Support**: Extract text from PDF and Word (.docx) files
    * **Progress Tracking**: Real-time progress monitoring for bulk operations
    * **Background Processing**: Asynchronous file processing with status updates
    * **Session Management**: Persistent sessions with configurable retention period
    * **Concurrent Processing**: Uses ProcessPoolExecutor for CPU-bound extraction tasks
    * **Production Ready**: Optimized for Railway deployment with proper resource management

    ### Supported File Types

    * **PDF**: Uses pdfplumber (pdfminer under the hood)
    * **Word**: Supports modern .docx files (not legacy .doc)

    ### Rate Limits

    * **File Size**: 20MB per file
    * **Files per Request**: 8 files maximum
    * **Concurrent Requests**: 3 maximum
    * **Session Retention**: 5 minutes after completion

    ### Authentication

    This API does not require authentication for basic usage.

    ### Error Handling

    The API provides comprehensive error handling with appropriate HTTP status codes:
    - `400`: Bad Request (invalid input)
    - `404`: Not Found (session not found)
    - `413`: Payload Too Large (file too large or too many files)
    - `429`: Too Many Requests (rate limit exceeded)
    - `500`: Internal Server Error (processing error)
    - `504`: Gateway Timeout (request timeout)
    """,
    version="1.1.0",
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json",
    contact={
        "name": "Document Text Extractor API",
        "url": "https://github.com/subhanshu/doc-to-text",
    },
    license_info={
        "name": "MIT",
    },
    servers=[
        {
            "url": "https://web-production-5932e.up.railway.app",
            "description": "Production server"
        },
        {
            "url": "http://localhost:8080",
            "description": "Development server"
        }
    ]
)

# Semaphore to limit concurrent requests
semaphore = asyncio.Semaphore(settings.MAX_CONCURRENT_REQUESTS)

# Process pool for CPU-bound extraction tasks
process_pool = ProcessPoolExecutor(max_workers=settings.PROCESS_POOL_WORKERS)

# ---- Utility functions ----

def save_upload_to_tempfile(upload: UploadFile, max_size: int) -> str:
    """Save UploadFile to a temp file and return path. Raises HTTPException for size limits."""
    suffix = os.path.splitext(upload.filename)[1]
    fd, path = tempfile.mkstemp(suffix=suffix, dir=settings.TEMP_DIR)
    os.close(fd)
    size = 0
    try:
        with open(path, "wb") as f:
            while True:
                chunk = upload.file.read(1024 * 32)
                if not chunk:
                    break
                size += len(chunk)
                if size > max_size:
                    raise HTTPException(status_code=413, detail=f"File too large: {upload.filename}")
                f.write(chunk)
        return path
    except Exception:
        try:
            os.remove(path)
        except Exception:
            pass
        raise


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        return "\n\n".join(texts)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Also try tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def detect_file_type_by_name(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".docm", ".dotx"}:
        return "docx"
    return "unknown"


def extract_single_file(path: str, filename: str, per_file_timeout: int) -> Tuple[str, str]:
    """Wrapper to detect type and call the correct extractor. Returns (filename, text) or raises."""
    ftype = detect_file_type_by_name(filename)
    if ftype == "pdf":
        text = extract_text_from_pdf(path)
        return filename, text
    elif ftype == "docx":
        text = extract_text_from_docx(path)
        return filename, text
    else:
        raise RuntimeError(f"Unsupported file type for file: {filename}")


async def run_extraction_in_pool(path: str, filename: str, timeout: int) -> Tuple[str, str]:
    loop = asyncio.get_running_loop()
    # Offload the cpu-bound extraction to worker processes
    future = loop.run_in_executor(process_pool, extract_single_file, path, filename, timeout)
    try:
        return await asyncio.wait_for(future, timeout=timeout)
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail=f"Extraction timed out for {filename}")
    except Exception as e:
        # If the work raised an exception, it will be wrapped here
        raise HTTPException(status_code=500, detail=f"Extraction failed for {filename}: {str(e)}")

# ---- Progress tracking functions ----

def update_progress(session_id: str, **kwargs):
    """Update progress for a session"""
    if session_id in progress_store:
        for key, value in kwargs.items():
            setattr(progress_store[session_id], key, value)
        logger.info(f"Progress updated for session {session_id}: {kwargs}")

def get_progress(session_id: str) -> ProgressStatus | None:
    """Get progress for a session"""
    return progress_store.get(session_id)

def cleanup_progress(session_id: str):
    """Clean up progress data after completion"""
    if session_id in progress_store:
        del progress_store[session_id]

def cleanup_expired_sessions():
    """Clean up sessions that have exceeded retention period"""
    current_time = datetime.now()
    expired_sessions = []
    
    for session_id, progress in progress_store.items():
        # Check if session is completed and has exceeded retention period
        if progress.end_time and progress.status in ['completed', 'failed']:
            time_since_completion = (current_time - progress.end_time).total_seconds()
            if time_since_completion > settings.SESSION_RETENTION_SECONDS:
                expired_sessions.append(session_id)
    
    # Clean up expired sessions
    for session_id in expired_sessions:
        del progress_store[session_id]
        logger.info(f"Cleaned up expired session: {session_id}")
    
    return len(expired_sessions)

def cleanup_oldest_sessions():
    """Clean up oldest completed sessions when we hit the limit"""
    if len(progress_store) <= settings.MAX_SESSIONS:
        return
    
    # Sort sessions by end_time (oldest first)
    completed_sessions = [
        (session_id, progress) 
        for session_id, progress in progress_store.items() 
        if progress.status in ['completed', 'failed'] and progress.end_time
    ]
    
    # Sort by end_time (oldest first)
    completed_sessions.sort(key=lambda x: x[1].end_time)
    
    # Remove oldest sessions until we're under the limit
    sessions_to_remove = len(progress_store) - settings.MAX_SESSIONS
    for i in range(min(sessions_to_remove, len(completed_sessions))):
        session_id = completed_sessions[i][0]
        del progress_store[session_id]
        logger.info(f"Cleaned up oldest session: {session_id}")

def schedule_session_cleanup():
    """Schedule periodic cleanup of expired sessions"""
    import threading
    import time
    
    def cleanup_worker():
        while True:
            try:
                time.sleep(60)  # Run cleanup every minute
                expired_count = cleanup_expired_sessions()
                if expired_count > 0:
                    logger.info(f"Cleaned up {expired_count} expired sessions")
            except Exception as e:
                logger.error(f"Error in session cleanup: {e}")
    
    # Start cleanup thread
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()
    logger.info("Session cleanup scheduler started")

# ---- API endpoints ----
@app.get("/", 
         summary="API Information",
         description="Get comprehensive information about the Document Text Extractor API including available endpoints, supported formats, and configuration limits.",
         response_model=APIInfoResponse,
         responses={
             200: {
                 "description": "API information retrieved successfully",
                 "content": {
                     "application/json": {
                         "example": {
                             "title": "Document Text Extractor API",
                             "version": "1.1.0",
                             "description": "A production-ready FastAPI service for extracting text from PDF and Word (.docx) documents with progress tracking",
                             "endpoints": {
                                 "health": "/health",
                                 "extract": "/extract",
                                 "extract_with_progress": "/extract-progress",
                                 "progress": "/progress/{session_id}",
                                 "sessions": "/sessions",
                                 "docs": "/docs",
                                 "redoc": "/redoc",
                                 "openapi": "/openapi.json"
                             },
                             "supported_formats": ["PDF (.pdf)", "Word (.docx, .docm, .dotx)"],
                             "max_file_size": "20MB",
                             "max_files_per_request": 8
                         }
                     }
                 }
             }
         },
         tags=["General"])
async def root():
    """
    Get API Information
    
    Returns comprehensive information about the Document Text Extractor API including:
    - API title, version, and description
    - Available endpoints and their purposes
    - Supported file formats
    - Configuration limits and constraints
    
    This endpoint is useful for:
    - API discovery and exploration
    - Understanding available functionality
    - Checking configuration limits
    - Integration planning
    """
    return APIInfoResponse(
        title="Document Text Extractor API",
        version="1.1.0",
        description="A production-ready FastAPI service for extracting text from PDF and Word (.docx) documents with progress tracking",
        endpoints={
            "health": "/health",
            "extract": "/extract",
            "extract_with_progress": "/extract-progress",
            "progress": "/progress/{session_id}",
            "sessions": "/sessions",
            "docs": "/docs",
            "redoc": "/redoc",
            "openapi": "/openapi.json"
        },
        supported_formats=["PDF (.pdf)", "Word (.docx, .docm, .dotx)"],
        max_file_size=f"{settings.MAX_FILE_SIZE_BYTES // (1024*1024)}MB",
        max_files_per_request=settings.MAX_TOTAL_FILES
    )


@app.get("/health", 
         summary="Health Check",
         description="Check the health status of the service and get configuration information",
         response_model=HealthResponse,
         responses={
             200: {
                 "description": "Service is healthy",
                 "content": {
                     "application/json": {
                         "example": {
                             "status": "ok",
                             "pool_workers": 3
                         }
                     }
                 }
             }
         },
         tags=["General"])
async def health():
    """
    Health Check
    
    Returns the current health status of the service and configuration information.
    
    This endpoint is useful for:
    - Monitoring service availability
    - Load balancer health checks
    - Debugging configuration issues
    - Checking process pool status
    
    Returns:
    - status: Service status ("ok" when healthy)
    - pool_workers: Number of process pool workers available
    """
    return HealthResponse(
        status="ok", 
        pool_workers=settings.PROCESS_POOL_WORKERS
    )


@app.post("/extract",
          summary="Extract Text from Documents",
          description="Extract text content from multiple PDF and Word (.docx) files in parallel and return all results at once",
          response_model=ExtractResponse,
          responses={
              200: {
                  "description": "Text extraction completed successfully",
                  "content": {
                      "application/json": {
                          "example": {
                              "results": [
                                  {
                                      "filename": "document1.pdf",
                                      "text": "Extracted text content from PDF..."
                                  },
                                  {
                                      "filename": "document2.docx",
                                      "text": "Extracted text content from Word document..."
                                  }
                              ],
                              "errors": []
                          }
                      }
                  }
              },
              400: {
                  "description": "Bad Request - No files uploaded",
                  "content": {
                      "application/json": {
                          "example": {"detail": "No files uploaded"}
                      }
                  }
              },
              413: {
                  "description": "Payload Too Large - Too many files or file too large",
                  "content": {
                      "application/json": {
                          "example": {"detail": "Too many files. Max allowed 8"}
                      }
                  }
              },
              429: {
                  "description": "Too Many Requests - Server is busy",
                  "content": {
                      "application/json": {
                          "example": {"detail": "Server is busy. Please try again later."}
                      }
                  }
              },
              504: {
                  "description": "Gateway Timeout - Request timed out",
                  "content": {
                      "application/json": {
                          "example": {"detail": "Request timed out"}
                      }
                  }
              }
          },
          tags=["Text Extraction"])
async def extract(files: List[UploadFile] = File(..., description="PDF or Word (.docx) files to extract text from")):
    """
    Extract Text from Documents
    
    Process multiple PDF and Word (.docx) files in parallel and return all extracted text at once.
    This endpoint is ideal for immediate processing when you need results quickly.
    
    **Supported File Types:**
    - PDF files (.pdf)
    - Word documents (.docx, .docm, .dotx)
    
    **File Limits:**
    - Maximum file size: 20MB per file
    - Maximum files per request: 8 files
    - Maximum concurrent requests: 3
    
    **Processing:**
    - Files are processed in parallel using ProcessPoolExecutor
    - CPU-bound extraction tasks are offloaded to worker processes
    - Results are returned only after all files are processed
    
    **Use Cases:**
    - Quick text extraction for small batches
    - Immediate processing when you need results right away
    - Simple integration without progress tracking
    
    **Example Usage:**
    ```bash
    curl -X POST "https://api.example.com/extract" \
         -H "Content-Type: multipart/form-data" \
         -F "files=@document1.pdf" \
         -F "files=@document2.docx"
    ```
    
    **Response:**
    - `results`: List of successfully processed files with extracted text
    - `errors`: List of any processing errors with error details
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    if len(files) > settings.MAX_TOTAL_FILES:
        raise HTTPException(status_code=413, detail=f"Too many files. Max allowed {settings.MAX_TOTAL_FILES}")

    # Acquire semaphore to limit concurrent requests
    try:
        await asyncio.wait_for(semaphore.acquire(), timeout=5)
    except asyncio.TimeoutError:
        raise HTTPException(status_code=429, detail="Server is busy. Try again later.")

    temp_paths = []
    results = []
    errors = []

    try:
        # Save all uploads to temp files first
        for upload in files:
            path = save_upload_to_tempfile(upload, settings.MAX_FILE_SIZE_BYTES)
            temp_paths.append((path, upload.filename))

        # Launch extraction tasks in parallel
        coros = [run_extraction_in_pool(path, fname, settings.PER_FILE_TIMEOUT) for path, fname in temp_paths]
        # Gather with overall timeout
        try:
            extracted = await asyncio.wait_for(asyncio.gather(*coros, return_exceptions=True), timeout=settings.REQUEST_TIMEOUT)
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Overall request timed out")

        for item in extracted:
            if isinstance(item, Exception):
                errors.append({"filename": "unknown", "error": str(item)})
            else:
                fname, text = item
                results.append({"filename": fname, "text": text})

        return JSONResponse(status_code=200, content={"results": results, "errors": errors})

    finally:
        # release semaphore and cleanup
        try:
            semaphore.release()
        except Exception:
            pass
        for path, _ in temp_paths:
            try:
                os.remove(path)
            except Exception:
                pass


@app.post("/extract-progress",
          summary="Extract Text with Progress Tracking",
          description="Extract text from files with real-time progress tracking using background processing",
          response_model=ExtractProgressResponse,
          responses={
              202: {
                  "description": "Processing started successfully",
                  "content": {
                      "application/json": {
                          "example": {
                              "session_id": "123e4567-e89b-12d3-a456-426614174000",
                              "status": "processing",
                              "message": "Processing started. Use /progress/{session_id} to track progress."
                          }
                      }
                  }
              },
              400: {
                  "description": "Bad Request - No files uploaded",
                  "content": {
                      "application/json": {
                          "example": {"detail": "No files uploaded"}
                      }
                  }
              },
              413: {
                  "description": "Payload Too Large - Too many files or file too large",
                  "content": {
                      "application/json": {
                          "example": {"detail": "Too many files. Max allowed 8"}
                      }
                  }
              }
          },
          tags=["Text Extraction", "Progress Tracking"])
async def extract_with_progress(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(..., description="PDF or Word (.docx) files to extract text from")
):
    """
    Extract Text with Progress Tracking
    
    Start background processing of multiple PDF and Word (.docx) files with real-time progress tracking.
    This endpoint is ideal for bulk operations where you need to monitor progress.
    
    **How it works:**
    1. Upload files to this endpoint
    2. Receive a session ID for tracking
    3. Use `/progress/{session_id}` to check progress
    4. Results are available when status is "completed"
    5. Session is retained for 5 minutes after completion
    
    **Supported File Types:**
    - PDF files (.pdf)
    - Word documents (.docx, .docm, .dotx)
    
    **File Limits:**
    - Maximum file size: 20MB per file
    - Maximum files per request: 8 files
    - Maximum concurrent requests: 3
    
    **Session Management:**
    - Sessions are retained for 5 minutes after completion
    - Use `/sessions` to list all active sessions
    - Use `/progress/{session_id}` to check progress
    - Use `DELETE /progress/{session_id}` to clean up early
    
    **Progress Tracking:**
    - Real-time progress percentage (0-100%)
    - Current file being processed
    - Elapsed time and estimated completion
    - Detailed results and error information
    
    **Example Usage:**
    ```bash
    # Start processing
    curl -X POST "https://api.example.com/extract-progress" \
         -H "Content-Type: multipart/form-data" \
         -F "files=@document1.pdf" \
         -F "files=@document2.docx"
    
    # Check progress (replace SESSION_ID)
    curl -X GET "https://api.example.com/progress/SESSION_ID"
    ```
    
    **Response:**
    - `session_id`: Unique identifier for tracking progress
    - `status`: Current processing status ("processing")
    - `message`: Instructions for tracking progress
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    if len(files) > settings.MAX_TOTAL_FILES:
        raise HTTPException(status_code=413, detail=f"Too many files. Max allowed {settings.MAX_TOTAL_FILES}")

    # Generate session ID
    session_id = str(uuid.uuid4())
    
    # Initialize progress
    progress_store[session_id] = ProgressStatus(
        session_id=session_id,
        status="processing",
        total_files=len(files),
        processed_files=0,
        start_time=datetime.now()
    )
    
    # Read file contents immediately to avoid I/O issues in background task
    file_data = []
    for upload in files:
        upload.file.seek(0)
        content = await upload.read()
        file_data.append({
            "filename": upload.filename,
            "content": content,
            "content_type": upload.content_type
        })
    
    # Start background processing
    background_tasks.add_task(process_files_with_progress, session_id, file_data)
    
    return JSONResponse(status_code=202, content={
        "session_id": session_id,
        "status": "processing",
        "message": "Processing started. Use /progress/{session_id} to track progress."
    })


async def process_files_with_progress(session_id: str, file_data: List[Dict[str, Any]]):
    """Background task to process files with progress updates"""
    
    # Acquire semaphore to limit concurrent requests
    try:
        await asyncio.wait_for(semaphore.acquire(), timeout=5)
    except asyncio.TimeoutError:
        update_progress(session_id, status="failed")
        return

    temp_paths = []
    results = []
    errors = []

    try:
        # Save all file data to temp files first
        for file_info in file_data:
            try:
                # Create temp file and write content
                suffix = os.path.splitext(file_info["filename"])[1]
                fd, path = tempfile.mkstemp(suffix=suffix, dir=settings.TEMP_DIR)
                os.close(fd)
                
                with open(path, "wb") as f:
                    f.write(file_info["content"])
                
                temp_paths.append((path, file_info["filename"]))
            except Exception as e:
                errors.append({"filename": file_info["filename"], "error": str(e)})

        # Process files one by one for better progress tracking
        for i, (path, filename) in enumerate(temp_paths):
            update_progress(
                session_id,
                processed_files=i,
                current_file=filename
            )
            
            try:
                fname, text = await run_extraction_in_pool(path, filename, settings.PER_FILE_TIMEOUT)
                results.append({"filename": fname, "text": text})
                logger.info(f"Successfully processed {filename} for session {session_id}")
            except Exception as e:
                error_msg = str(e)
                errors.append({"filename": filename, "error": error_msg})
                logger.error(f"Failed to process {filename} for session {session_id}: {error_msg}")

        # Final update
        update_progress(
            session_id,
            status="completed",
            processed_files=len(temp_paths),
            current_file=None,
            results=results,
            errors=errors,
            end_time=datetime.now()
        )
        
        logger.info(f"Completed processing session {session_id}: {len(results)} successful, {len(errors)} failed")

    except Exception as e:
        logger.error(f"Fatal error in session {session_id}: {str(e)}")
        update_progress(
            session_id,
            status="failed",
            errors=[{"filename": "system", "error": str(e)}],
            end_time=datetime.now()
        )
    finally:
        # release semaphore and cleanup temp files
        try:
            semaphore.release()
        except Exception:
            pass
        for path, _ in temp_paths:
            try:
                os.remove(path)
            except Exception:
                pass


@app.get("/progress/{session_id}",
         summary="Get Processing Progress",
         description="Get real-time progress for a bulk processing session")
async def get_processing_progress(session_id: str):
    """Get progress for a processing session"""
    # Clean up expired sessions before checking
    cleanup_expired_sessions()
    
    progress = get_progress(session_id)
    
    if not progress:
        raise HTTPException(status_code=404, detail="Session not found or expired")
    
    # Calculate elapsed time
    elapsed_time = (
        (progress.end_time or datetime.now()) - progress.start_time
    ).total_seconds()
    
    # Calculate progress percentage
    progress_percentage = (
        (progress.processed_files / progress.total_files * 100) 
        if progress.total_files > 0 else 0
    )
    
    # Calculate time until expiration for completed sessions
    time_until_expiry = None
    if progress.end_time and progress.status in ['completed', 'failed']:
        time_since_completion = (datetime.now() - progress.end_time).total_seconds()
        time_until_expiry = max(0, settings.SESSION_RETENTION_SECONDS - time_since_completion)
    
    return {
        "session_id": progress.session_id,
        "status": progress.status,
        "progress": round(progress_percentage, 1),
        "total_files": progress.total_files,
        "processed_files": progress.processed_files,
        "current_file": progress.current_file,
        "results": progress.results,
        "errors": progress.errors,
        "elapsed_time": round(elapsed_time, 1),
        "start_time": progress.start_time.isoformat(),
        "end_time": progress.end_time.isoformat() if progress.end_time else None,
        "time_until_expiry": round(time_until_expiry, 1) if time_until_expiry is not None else None
    }


@app.get("/sessions",
         summary="List Active Sessions",
         description="Get list of all active sessions with their status")
async def list_sessions():
    """List all active sessions"""
    # Clean up expired sessions first
    cleanup_expired_sessions()
    
    sessions = []
    for session_id, progress in progress_store.items():
        elapsed_time = (
            (progress.end_time or datetime.now()) - progress.start_time
        ).total_seconds()
        
        progress_percentage = (
            (progress.processed_files / progress.total_files * 100) 
            if progress.total_files > 0 else 0
        )
        
        time_until_expiry = None
        if progress.end_time and progress.status in ['completed', 'failed']:
            time_since_completion = (datetime.now() - progress.end_time).total_seconds()
            time_until_expiry = max(0, settings.SESSION_RETENTION_SECONDS - time_since_completion)
        
        sessions.append({
            "session_id": session_id,
            "status": progress.status,
            "progress": round(progress_percentage, 1),
            "total_files": progress.total_files,
            "processed_files": progress.processed_files,
            "current_file": progress.current_file,
            "elapsed_time": round(elapsed_time, 1),
            "start_time": progress.start_time.isoformat(),
            "end_time": progress.end_time.isoformat() if progress.end_time else None,
            "time_until_expiry": round(time_until_expiry, 1) if time_until_expiry is not None else None
        })
    
    return {
        "total_sessions": len(sessions),
        "sessions": sessions
    }


@app.delete("/progress/{session_id}",
           summary="Clean Up Session",
           description="Clean up progress data for a completed session")
async def cleanup_session(session_id: str):
    """Clean up progress data for a session"""
    if session_id not in progress_store:
        raise HTTPException(status_code=404, detail="Session not found")
    
    cleanup_progress(session_id)
    return {"message": f"Session {session_id} cleaned up"}


# ---- Startup and shutdown events ----
@app.on_event("startup")
def startup_event():
    """Initialize resources on startup"""
    logger.info("Starting up Document Text Extractor API...")
    # Start session cleanup scheduler
    schedule_session_cleanup()
    logger.info("Startup complete")

@app.on_event("shutdown")
def shutdown_event():
    """Clean up resources on shutdown"""
    logger.info("Shutting down...")
    process_pool.shutdown(wait=True)
    logger.info("Shutdown complete")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
