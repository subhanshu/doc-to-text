"""
Production-ready FastAPI service: PDF & Word (.docx) text extraction
- Accepts multiple files per request
- Offloads CPU-bound extraction to a ProcessPoolExecutor (uses multiple cores)
- Limits concurrency with a semaphore to protect a 4-core / 8GB VM
- Robust error handling, timeouts, file-size limits, logging

Files: main.py (this file)
Usage (dev):
  pip install -r requirements.txt
  uvicorn main:app --host 0.0.0.0 --port 8080 --workers 1

Production tips (on 4-core VM):
  - Run with gunicorn + uvicorn workers, e.g.:
      gunicorn -k uvicorn.workers.UvicornWorker -w 3 main:app
    (3 workers recommended on a 4-core VM; each worker uses the ProcessPoolExecutor
     for per-request parallelism)
  - Optionally containerize with included Dockerfile snippet (see README comments)

Notes:
  - PDF extraction uses pdfplumber (pdfminer under the hood). It's fairly reliable
    for text-based PDFs. For scanned PDFs, enable optional OCR fallback (requires
    tesseract + poppler and extra python deps); instructions included in comments.
  - Word extraction supports modern .docx files. Older .doc (binary) is not supported
    by default; if needed add textract or antiword and handle as an optional plugin.

"""

import os
import shutil
import tempfile
import logging
import asyncio
from concurrent.futures import ProcessPoolExecutor, TimeoutError as FuturesTimeout
from typing import List, Tuple

from fastapi import FastAPI, File, UploadFile, HTTPException, status
from fastapi.responses import JSONResponse
from pydantic import BaseSettings

# Extraction libraries
import pdfplumber
import docx

# ---- Settings ----
class Settings(BaseSettings):
    MAX_FILE_SIZE_BYTES: int = 20 * 1024 * 1024  # 20 MB per file
    MAX_TOTAL_FILES: int = 8
    # Number of processes for ProcessPoolExecutor. Keep <= number of CPU cores.
    PROCESS_POOL_WORKERS: int = max(1, (os.cpu_count() or 1) - 1)
    # Limit concurrent requests to avoid OOM/CPU overcommit on small VM.
    MAX_CONCURRENT_REQUESTS: int = 3
    # Per-file extraction timeout (seconds)
    PER_FILE_TIMEOUT: int = 30
    # Overall request timeout
    REQUEST_TIMEOUT: int = 120
    # Temp dir base
    TEMP_DIR: str = None

    class Config:
        env_prefix = "EXTRACT_"

settings = Settings()

# Fix sensible defaults for a 4-core, 8GB VM if automatic detection returned large counts
if (os.cpu_count() or 1) >= 4:
    # leave one core for the event loop and webserver
    settings.PROCESS_POOL_WORKERS = min(settings.PROCESS_POOL_WORKERS, 3)
    settings.MAX_CONCURRENT_REQUESTS = min(settings.MAX_CONCURRENT_REQUESTS, 3)

# ---- Logging ----
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("extractor")

# ---- FastAPI app ----
app = FastAPI(
    title="Document Text Extractor API",
    description="A production-ready FastAPI service for extracting text from PDF and Word (.docx) documents",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json"
)

# Semaphore to limit concurrent requests
semaphore = asyncio.Semaphore(settings.MAX_CONCURRENT_REQUESTS)

# Process pool for CPU-bound extraction tasks
process_pool = ProcessPoolExecutor(max_workers=settings.PROCESS_POOL_WORKERS)

# ---- Utility functions ----

def save_upload_to_tempfile(upload: UploadFile, max_size: int) -> str:
    """Save UploadFile to a temp file and return path. Raises HTTPException for size limits."""
    suffix = os.path.splitext(upload.filename)[1]
    fd, path = tempfile.mkstemp(suffix=suffix, dir=settings.TEMP_DIR)
    os.close(fd)
    size = 0
    try:
        with open(path, "wb") as f:
            while True:
                chunk = upload.file.read(1024 * 32)
                if not chunk:
                    break
                size += len(chunk)
                if size > max_size:
                    raise HTTPException(status_code=413, detail=f"File too large: {upload.filename}")
                f.write(chunk)
        return path
    except Exception:
        try:
            os.remove(path)
        except Exception:
            pass
        raise


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF file using pdfplumber.
    This function runs in a separate process (no FastAPI objects here).
    """
    try:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        return "\n\n".join(texts)
    except Exception as e:
        # Bubble up a helpful message to the caller
        raise RuntimeError(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(path: str) -> str:
    """Extract text from a .docx file using python-docx.
    Runs in separate process.
    """
    try:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Also try tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {str(e)}")


def detect_file_type_by_name(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".docm", ".dotx"}:
        return "docx"
    return "unknown"


def extract_single_file(path: str, filename: str, per_file_timeout: int) -> Tuple[str, str]:
    """Wrapper to detect type and call the correct extractor. Returns (filename, text) or raises."""
    ftype = detect_file_type_by_name(filename)
    if ftype == "pdf":
        text = extract_text_from_pdf(path)
        return filename, text
    elif ftype == "docx":
        text = extract_text_from_docx(path)
        return filename, text
    else:
        raise RuntimeError(f"Unsupported file type for file: {filename}")


async def run_extraction_in_pool(path: str, filename: str, timeout: int) -> Tuple[str, str]:
    loop = asyncio.get_running_loop()
    # Offload the cpu-bound extraction to worker processes
    future = loop.run_in_executor(process_pool, extract_single_file, path, filename, timeout)
    try:
        return await asyncio.wait_for(future, timeout=timeout)
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail=f"Extraction timed out for {filename}")
    except Exception as e:
        # If the work raised an exception, it will be wrapped here
        raise HTTPException(status_code=500, detail=f"Extraction failed for {filename}: {str(e)}")


# ---- API endpoints ----
@app.get("/health", 
         summary="Health Check",
         description="Check the health status of the service and get configuration information",
         response_description="Service status and configuration details")
async def health():
    """
    Health check endpoint that returns the service status and configuration.
    
    Returns:
        dict: Service status and process pool worker count
    """
    return {"status": "ok", "pool_workers": settings.PROCESS_POOL_WORKERS}


@app.post("/extract")
async def extract(files: List[UploadFile] = File(...)):
    """Extract text from multiple uploaded files in parallel.

    Response: {
      "results": [ {"filename": str, "text": str }, ... ],
      "errors": [ {"filename": str, "error": str}, ... ]
    }
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    if len(files) > settings.MAX_TOTAL_FILES:
        raise HTTPException(status_code=413, detail=f"Too many files. Max allowed {settings.MAX_TOTAL_FILES}")

    # Acquire semaphore to limit concurrent requests
    try:
        await asyncio.wait_for(semaphore.acquire(), timeout=5)
    except asyncio.TimeoutError:
        raise HTTPException(status_code=429, detail="Server is busy. Try again later.")

    temp_paths = []
    tasks = []
    results = []
    errors = []

    try:
        # Save all uploads to temp files first (this is I/O bound and done serially to avoid inflating memory)
        for upload in files:
            if upload.content_type and upload.content_type.startswith("text/"):
                # allow small text files but still enforce size
                pass
            path = save_upload_to_tempfile(upload, settings.MAX_FILE_SIZE_BYTES)
            temp_paths.append((path, upload.filename))

        # Launch extraction tasks in parallel (but limited by process pool size)
        coros = [run_extraction_in_pool(path, fname, settings.PER_FILE_TIMEOUT) for path, fname in temp_paths]
        # Gather with overall timeout
        try:
            extracted = await asyncio.wait_for(asyncio.gather(*coros, return_exceptions=True), timeout=settings.REQUEST_TIMEOUT)
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Overall request timed out")

        for item in extracted:
            if isinstance(item, Exception):
                # extract filename if possible from exception message
                errors.append({"filename": "unknown", "error": str(item)})
            else:
                fname, text = item
                results.append({"filename": fname, "text": text})

        return JSONResponse(status_code=200, content={"results": results, "errors": errors})

    finally:
        # release semaphore and cleanup
        try:
            semaphore.release()
        except Exception:
            pass
        for path, _ in temp_paths:
            try:
                os.remove(path)
            except Exception:
                pass


# ---- Graceful shutdown ----
@app.on_event("shutdown")
def shutdown_event():
    logger.info("Shutting down process pool")
    process_pool.shutdown(wait=True)


# ---- Example Dockerfile (comment) ----
DOCKERFILE = r"""
# Example Dockerfile
FROM python:3.11-slim

# Install system deps for pdfplumber/pdfminer if needed
RUN apt-get update && apt-get install -y --no-install-recommends \
    build-essential \
    poppler-utils \
    && rm -rf /var/lib/apt/lists/*

WORKDIR /app
COPY . /app
RUN pip install -r requirements.txt

# Use gunicorn in production
CMD ["gunicorn", "-k", "uvicorn.workers.UvicornWorker", "-w", "3", "main:app", "--bind", "0.0.0.0:8080"]
"""


# ---- requirements.txt (comment) ----
REQUIREMENTS = r"""
fastapi>=0.95
uvicorn[standard]>=0.22
gunicorn>=20.1
pdfplumber>=0.7.5
python-docx>=0.8.11
pydantic>=1.10
"""

# ---- Optional OCR fallback notes ----
OCR_NOTES = r"""
If you need to handle scanned PDFs (images):

System packages to install (Debian/Ubuntu):
  apt-get install -y tesseract-ocr poppler-utils libpoppler-cpp-dev

Python packages:
  pip install pdf2image pillow pytesseract

Strategy:
  - Use pdf2image to render PDF pages to images, then run pytesseract.image_to_string on each page.
  - Because OCR is CPU and memory intensive, keep this disabled by default and enable
    behind a feature flag or separate endpoint.
"""

# End of file
